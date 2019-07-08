from bs4 import BeautifulSoup as soup
from selenium import webdriver
import requests
import re
import os

prices = [(0.7,0.8), (1,1.1), (1.5,1.6), (2, 2.1), (2.4, 2.55), (3.1, 3.2), (3.85, 3.95), (4.65, 4.75), (6.2,6.3), (7.8 ,7.9), (10.9, 11), (15.6, 15.7), (20.3, 20.4), (31.25, 31.35), (46.95, 47.05)]
urls = ['https://www.amazon.co.uk/b/ref=s9_acss_bw_cg_alukcat_2a1_w?node=7424489031&pf_rd_s=merchandised-search-12',
        'https://www.amazon.co.uk/b/ref=s9_acss_bw_cg_alukcat_2c1_w?node=7424498031&pf_rd_s=merchandised-search-12',
        'https://www.amazon.co.uk/b/ref=s9_acss_bw_cg_alukcat_2d1_w?node=7424508031&pf_rd_s=merchandised-search-12',
        'https://www.amazon.co.uk/b/ref=s9_acss_bw_cg_alukcat_2e1_w?node=7424514031&pf_rd_s=merchandised-search-12',
        'https://www.amazon.co.uk/b/ref=alukgwhc?node=7212961031&pf_rd_p=b1de4cf9-af32-452b-9dc5-ae1eab19caef&pf_rd_r=A8HZSRCTY15JSS6BX705'] 

        #URLS[4] has different layout

products = []

def sort_products(prd):
    combined = [j for i in prd for j in i]
    combined.sort(key=lambda x: float(x['price']))
    return combined

def get_results(url):
    results = []
    driver = webdriver.Chrome()
    driver.get(url)
    

    for a in range(3):
        for b in range(3):
            cur = [x for x in driver.find_elements_by_class_name('a-carousel-card')]
            cur2 = []
            img_tag = ''
            for c in cur:
                s = soup(c.get_attribute('innerHTML'), 'html.parser')
                try:
                    img_tag = s.select('img.aok-align-center')[0]['src']
                except:
                    pass
                cur2.append(c.text + '\n' + img_tag)
            results += cur2
            driver.find_elements_by_class_name('a-button-inner')[a].click()

    length = len(results)
    for x in range(length-1, 0, -1):
        if not results[x]:
            del results[x]
            length-=1
    #print([result.split('\n') for result in results])
    #print('*'*100)
    #print('UNCLEANED RESULTS: ', [result.split('\n') for result in results])
    for x in range(len(results)):
        split = results[x].split('\n')
        obj = {}
        obj['title'] = split[0]
        for s in split:
            if 'http' in s:
                obj['img'] = s
            for t in s.split(' '):
                if '£' in t:
                    obj['price'] = float(t.replace('(', '').replace(')', '').replace('£', ''))
        results[x] = obj
        
    # CHECK IF THERE IS A PRICE SOME ITEMS ARE UNAVAILABLE
    for x in range(len(results)-1, 0, -1):
        if 'price' not in results[x]:
            del results[x]

    #results.sort(key=lambda x: float(x['price']))
    #print(results)
    #print('*'*100)
    results = [dict(t) for t in {tuple(d.items()) for d in results}]
    results.sort(key=lambda x: x['price'])
    print(results)
    driver.close()
    return results

def get_results_two(url):
    results = []
    driver = webdriver.Chrome()
    driver.get(url)

    results += [x.text for x in driver.find_elements_by_class_name('s-result-item')]
    driver.find_element_by_class_name('pagnNextArrow').click()
    for x in range(5):
        cur = [x for x in driver.find_elements_by_class_name('s-result-item')]
        cur2 = []
        img_tag = ''
        for c in cur:
            s = soup(c.get_attribute('innerHTML'), 'html.parser')
            try:
                img_tag = s.select('img.s-image')[0]['src']
            except:
                pass
            cur2.append(c.text + '\n' + img_tag)

        results += cur2
        driver.find_element_by_class_name('a-last').click()

    length = len(results)
    for x in range(length-1, 0, -1):
        if not results[x]:
            del results[x]
            length-=1
    #print([result.split('\n') for result in results])
    #print('*'*100)
    #print('UNCLEANED RESULTS: ', [result.split('\n') for result in results])
    for x in range(len(results)):
        split = results[x].split('\n')
        obj = {}
        obj['title'] = split[0]
        for s in split:
            if 'http' in s:
                obj['img'] = s 
            for t in s.split(' '):
                if '£' in t:
                    try:
                        obj['price'] = float(t.split('£')[1])
                    except:
                        pass
        results[x] = obj

    # CHECK IF THERE IS A PRICE SOME ITEMS ARE UNAVAILABLE
    for x in range(len(results)-1, 0, -1):
        if 'price' not in results[x]:
            del results[x]

    #results.sort(key=lambda x: float(x['price']))
    #print(results)
    #print('*'*100)
    results = [dict(t) for t in {tuple(d.items()) for d in results}]
    results.sort(key=lambda x: x['price'])

    #print('RESUTLS', results)
    return results


products.append(get_results(urls[0]))
products.append(get_results(urls[1]))
products.append(get_results(urls[2]))
products.append(get_results(urls[3]))
products.append(get_results_two(urls[4]))
products = sort_products(products)
#print(products)

final_products = []
for product in products:
    for price in prices:
        if product['price'] > price[0] and product['price'] < price[1]:
            final_products.append(product)

print('FINAL PRODUCTS: ', final_products)

def download_image(url):
    f = open('tmp/'+os.path.basename(url), 'wb')
    f.write(requests.get(url).content)
    f.close()


from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Products', 0)

for product in final_products:
    download_image(product['img'])
    document.add_paragraph('Description: {}'.format(product['title']))
    document.add_paragraph('Price: {}'.format(product['price']))
    document.add_picture('tmp/'+os.path.basename(product['img']))

document.save('products.docx')